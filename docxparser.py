from docx import Document

def read_docx_with_tables(file_path):
    doc = Document(file_path)
    content = []

    # Read all paragraphs
    for para in doc.paragraphs:
        if para.text.strip():  # Skip empty paragraphs
            content.append(para.text)

    # Read all tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text.strip())
            content.append('\t'.join(row_text))  # Use tab or comma to separate cells

    return '\n'.join(content)

# Example usage
if __name__ == "__main__":
    path = "SHA Draft-Investor friendly(2).docx"  # Replace with your actual file path
    full_text = read_docx_with_tables(path)
    print("Full Document Content (including tables):\n")
    print(len(full_text))
